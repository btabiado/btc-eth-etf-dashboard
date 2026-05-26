"""Generate a Word document inventorying dashboard data sources by tab (v2)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# (Tab name, description, [(provider, endpoint/domain, what it feeds)])
TABS = [
    ("AI News", "AI/tech headlines, funding, and YC company tracking", [
        ("Hacker News (Algolia)", "hn.algolia.com/api/v1/search", "AI story search + link resolution"),
        ("TechCrunch", "techcrunch.com/feed/ (+ AI category)", "AI news RSS"),
        ("VentureBeat", "venturebeat.com/feed/ (+ AI category)", "AI/tech news RSS"),
        ("OpenAI", "openai.com/news/rss.xml", "Lab announcements RSS"),
        ("Anthropic", "anthropic.com/news/feed.xml", "Lab announcements RSS"),
        ("Ars Technica", "feeds.arstechnica.com/arstechnica/index/", "Tech news RSS"),
        ("The Verge", "theverge.com/ai-artificial-intelligence/rss/index.xml", "AI news RSS"),
        ("MIT Technology Review", "technologyreview.com/.../artificial-intelligence/feed", "AI research RSS"),
        ("Y Combinator", "yc-oss.github.io/api/tags/artificial-intelligence.json", "YC AI companies"),
        ("SEC EDGAR (Form D)", "efts.sec.gov/LATEST/search-index", "Recent AI private placements"),
        ("Wikipedia", "Infobox API/scrape", "Company metadata enrichment"),
    ]),
    ("Crypto", "Spot prices, market caps, volume, Fear & Greed", [
        ("CoinGecko", "api.coingecko.com/api/v3", "Price, volume, market cap, trending, global stats"),
        ("Coinbase Exchange", "api.exchange.coinbase.com", "Bid/ask, 24h range/volume"),
        ("CoinDesk", "data-api.coindesk.com", "CADLI daily OHLC index"),
        ("Alternative.me", "api.alternative.me/fng/", "Fear & Greed index"),
        ("CryptoCompare", "min-api.cryptocompare.com / data-api.cryptocompare.com", "Historical OHLC, social/news depth"),
    ]),
    ("Crypto Signals", "Derived momentum/volatility signals", [
        ("CoinGecko", "api.coingecko.com/api/v3", "Underlying price/volume series"),
        ("Deribit", "deribit.com/api/v2 (DVOL)", "BTC/ETH implied volatility"),
        ("Alternative.me", "api.alternative.me/fng/", "Sentiment input"),
    ]),
    ("Whale Activity", "Large on-chain transactions across chains", [
        ("blockchain.info", "api.blockchain.info/charts", "BTC on-chain metrics"),
        ("Blockchair", "api.blockchair.com (ETH, LTC, BCH, DOGE)", "Whale tx, multi-chain stats"),
        ("Etherscan v2", "api.etherscan.io/v2/api", "ETH gas, burn, largest tx, supply"),
        ("mempool.space", "mempool.space/api", "BTC fees, hashrate, whale tx scan"),
        ("Coin Metrics", "community-api.coinmetrics.io/v4", "ETH whale/address series"),
    ]),
    ("Point of Control", "Volume-profile / POC analysis", [
        ("CoinGecko", "api.coingecko.com/api/v3", "Price/volume for volume profile"),
        ("Coinbase Exchange", "api.exchange.coinbase.com", "Range/volume inputs"),
    ]),
    ("Research", "Social depth + crypto news sentiment", [
        ("CryptoCompare", "min-api.cryptocompare.com/data/social", "Social depth, news sentiment"),
        ("Santiment", "api.santiment.net/graphql", "Daily active addresses (optional)"),
        ("Reddit", "oauth.reddit.com / reddit.com/r/*/top/.rss", "Subreddit subscriber/stats"),
        ("CoinDesk", "coindesk.com/.../rss", "Crypto news RSS"),
        ("Cointelegraph", "cointelegraph.com/rss", "Crypto news RSS"),
        ("Decrypt", "decrypt.co/feed", "Crypto news RSS"),
        ("Bitcoin Magazine", "bitcoinmagazine.com/feed", "Bitcoin news RSS"),
        ("The Block", "theblock.co/rss.xml", "Crypto research RSS"),
    ]),
    ("DeFi", "TVL, DEX volume, fees, yields", [
        ("DeFiLlama", "api.llama.fi / yields.llama.fi / stablecoins.llama.fi", "TVL, DEX vol, fees, stablecoin yields, bridges"),
        ("GeckoTerminal", "api.geckoterminal.com/api/v2", "Trending/new DEX pools"),
    ]),
    ("ETF Flows", "BTC/ETH spot ETF daily net flows", [
        ("Farside Investors", "farside.co.uk/bitcoin-etf-flow-all-data/", "Daily flows (manual paste workflow)"),
        ("canadiancode (GitHub mirror)", "raw.githubusercontent.com/canadiancode/btc-etf-flows", "Community BTC ETF flows"),
        ("SoSoValue", "api.sosovalue.com/openapi/v2", "Historical net inflows (paid)"),
        ("CoinGlass v4", "coinglass ETF-flow endpoint", "ETF flows (paid)"),
    ]),
    ("Futures", "Perp funding, open interest, long/short", [
        ("Coinbase International", "api.international.coinbase.com/api/v1", "Funding, mark price, OI, volume"),
        ("OKX", "okx.com/api/v5", "Funding history, long/short ratio, OI"),
        ("Deribit", "deribit.com/api/v2 (DVOL)", "Implied volatility"),
    ]),
    ("Stocks", "Equity prices, news sentiment, fundamentals", [
        ("Twelvedata (NEW in v2)", "api.twelvedata.com/time_series", "Client-side stock prices (browser key)"),
        ("Alpha Vantage", "alphavantage.co/query", "Client-side fallback prices + news sentiment"),
        ("Yahoo Finance", "query1.finance.yahoo.com", "Most-active list, OHLCV, hourly"),
        ("Finnhub", "finnhub.io/api/v1", "Per-ticker news/sentiment, analyst recs, earnings"),
        ("SEC EDGAR", "data.sec.gov / sec.gov/Archives", "Company facts, 13F, Form 4, 8-K"),
    ]),
    ("Travel Advisories", "Per-country risk levels & alerts", [
        ("U.S. State Department", "travel.state.gov (HTML + TAsTWs.xml RSS)", "Advisory levels 1-4, risk codes, alerts"),
    ]),
    ("CPI", "Inflation series", [
        ("FRED (St. Louis Fed)", "api.stlouisfed.org/fred/series/observations", "CPI/PCE series (20+ across 7 categories)"),
    ]),
    ("Supplies", "Supply-chain pressure & port throughput", [
        ("NY Fed", "newyorkfed.org (GSCPI)", "Global Supply Chain Pressure Index"),
        ("LA County Open Data", "data.lacity.org/resource/tsuv-4rgh.json", "Port of LA monthly TEU"),
        ("NY State Open Data", "data.ny.gov/resource/629s-5a55.json", "Port of NY/NJ monthly TEU"),
        ("EIA", "api.eia.gov/v2", "WTI/Brent/gasoline (energy)"),
    ]),
    ("Metals", "Gold & silver prices, holdings, production", [
        ("FRED", "api.stlouisfed.org (GOLDPMGBD228NLBM)", "London Gold PM fix"),
        ("Yahoo Finance", "query1.finance.yahoo.com (GC=F, SI=F)", "Gold/silver futures"),
        ("IMF SDMX", "api.imf.org/external/sdmx/2.1", "Central-bank gold holdings"),
        ("USGS ScienceBase", "sciencebase.gov/catalog", "World mine production by country"),
    ]),
    ("UAP", "UFO sighting records + official references", [
        ("NUFORC", "nuforc.org/subndx + admin-ajax.php", "Live sightings (2014+)"),
        ("planetsig (GitHub)", "raw.githubusercontent.com/planetsig/ufo-reports", "Historical sightings 1906-2014"),
        ("timothyrenner (GitHub)", "raw.githubusercontent.com/timothyrenner/nuforc_sightings_data", "Alternate processed NUFORC data"),
        ("Dept. of War — PURSUE portal", "war.gov/ufo/ (reference)", "Declassified UAP file releases"),
        ("AARO", "aaro.mil (reference)", "All-domain Anomaly Resolution Office"),
        ("NARA — Project Blue Book", "archives.gov (reference)", "Historical USAF UFO records"),
        ("CIA Reading Room", "cia.gov/readingroom (reference)", "Declassified UFO collection"),
        ("FBI Vault", "vault.fbi.gov/UFO (reference)", "FBI UFO files"),
        ("ODNI", "dni.gov (reference)", "2024 Consolidated UAP report"),
    ]),
]

CROSS = [
    ("FRED", "api.stlouisfed.org", "Macro overlays (rates, DXY, M2) across multiple tabs"),
    ("Anthropic Claude API", "api.anthropic.com/v1/messages", "AI chat / narratives (key can be set client-side in v2)"),
]

def add_source_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Provider", "Endpoint / Domain", "Feeds")):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for prov, ep, feeds in rows:
        c = table.add_row().cells
        c[0].text, c[1].text, c[2].text = prov, ep, feeds
        for cell in c:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

doc = Document()

title = doc.add_heading("BTC/ETH ETF Dashboard — Data Sources by Tab", level=0)

unique = set()
total_refs = 0
for _, _, rows in TABS:
    for prov, _, _ in rows:
        unique.add(prov.split(" (")[0].replace(" (NEW in v2)", "").replace(" — PURSUE portal", "").strip())
        total_refs += 1
for prov, _, _ in CROSS:
    unique.add(prov)

intro = doc.add_paragraph()
intro.add_run("Generated for the v2 dashboard. ").italic = True
intro.add_run(
    f"The dashboard has {len(TABS)} tabs drawing on {len(unique)} unique data providers "
    f"({total_refs} tab-level source references, since several providers feed multiple tabs). "
    "Most feeds are fetched server-side by the fetch_*.py builders and written to JSON sidecars "
    "the static site loads; Stocks (Twelvedata/Alpha Vantage) and the AI chat are fetched "
    "client-side with user-supplied keys."
).font.size = Pt(10)

doc.add_heading("Summary", level=1)
s = doc.add_paragraph(style="List Bullet")
s.add_run(f"Dashboard tabs: {len(TABS)}")
s = doc.add_paragraph(style="List Bullet")
s.add_run(f"Unique data providers: {len(unique)}")
s = doc.add_paragraph(style="List Bullet")
s.add_run(f"Total tab-level source references: {total_refs}")
s = doc.add_paragraph(style="List Bullet")
s.add_run(f"Cross-cutting sources (used across tabs): {len(CROSS)}")

for name, desc, rows in TABS:
    doc.add_heading(f"{name}  ({len(rows)} sources)", level=1)
    p = doc.add_paragraph()
    p.add_run(desc).italic = True
    add_source_table(doc, rows)

doc.add_heading("Cross-cutting sources", level=1)
doc.add_paragraph("Used across multiple tabs rather than tied to one.").runs[0].italic = True
add_source_table(doc, CROSS)

import os
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "v2", "Dashboard_Data_Sources.docx")
doc.save(out)
print("Saved", out)
print("Tabs:", len(TABS), "Unique providers:", len(unique), "Total refs:", total_refs)
